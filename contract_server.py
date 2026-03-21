"""
Meezan Contract Generation Server
Generates real Word (.docx) contracts using OpenAI GPT-4.1-mini + python-docx
Runs on port 5051
"""
import os, io, json, re
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

app = Flask(__name__)
CORS(app)
client = OpenAI()

# ── Contract type → system prompt + required fields ──────────────────────────
CONTRACT_CONFIGS = {
    "عقد إيجار تجاري": {
        "system": """أنت محامٍ متخصص في عقود الإيجار التجاري. اكتب عقد إيجار تجاري كامل ومفصل باللغة العربية الفصحى وفق القانون المحدد. 
العقد يجب أن يتضمن: ديباجة، تعريفات، موضوع العقد، مدة الإيجار، بدل الإيجار وطريقة السداد، الضمان، التزامات المؤجر، التزامات المستأجر، الصيانة، التأمين، إنهاء العقد، فض النزاعات، أحكام ختامية، توقيعات.
استخدم صياغة قانونية دقيقة ومواد مرقمة. أعد النص فقط بدون أي تعليق.""",
        "fields": ["العقار وعنوانه", "المساحة (م²)", "بدل الإيجار الشهري", "مدة الإيجار", "قيمة الضمان البنكي", "الغرض من الاستخدام"]
    },
    "عقد بيع عقار": {
        "system": """أنت محامٍ متخصص في عقود بيع العقارات. اكتب عقد بيع عقار كامل ومفصل باللغة العربية الفصحى وفق القانون المحدد.
العقد يجب أن يتضمن: ديباجة، وصف العقار، ثمن البيع، طريقة الدفع، تاريخ التسليم، ضمانات البائع، التزامات المشتري، نقل الملكية، الرسوم والضرائب، فض النزاعات، أحكام ختامية، توقيعات.
استخدم صياغة قانونية دقيقة ومواد مرقمة. أعد النص فقط بدون أي تعليق.""",
        "fields": ["وصف العقار ورقم الصك", "ثمن البيع الإجمالي", "طريقة الدفع (نقد/أقساط)", "تاريخ التسليم", "الرسوم والضرائب (على من تقع)"]
    },
    "عقد شراكة": {
        "system": """أنت محامٍ متخصص في عقود الشراكة التجارية. اكتب عقد شراكة كامل ومفصل باللغة العربية الفصحى وفق القانون المحدد.
العقد يجب أن يتضمن: ديباجة، اسم الشركة ومجال عملها، رأس المال وحصص الشركاء، إدارة الشركة، توزيع الأرباح والخسائر، اتخاذ القرارات، حظر المنافسة، انسحاب الشريك، حل الشراكة، فض النزاعات، أحكام ختامية، توقيعات.
استخدم صياغة قانونية دقيقة ومواد مرقمة. أعد النص فقط بدون أي تعليق.""",
        "fields": ["اسم الشركة ومجال العمل", "رأس المال الإجمالي", "حصة الطرف الأول (%)", "حصة الطرف الثاني (%)", "نسبة توزيع الأرباح", "مدة الشراكة"]
    },
    "عقد توريد": {
        "system": """أنت محامٍ متخصص في عقود التوريد والمشتريات. اكتب عقد توريد كامل ومفصل باللغة العربية الفصحى وفق القانون المحدد.
العقد يجب أن يتضمن: ديباجة، موضوع التوريد، الكميات والمواصفات، الأسعار وشروط الدفع، مواعيد التسليم، ضمان الجودة، التفتيش والقبول، العقوبات التأخيرية، القوة القاهرة، إنهاء العقد، فض النزاعات، أحكام ختامية، توقيعات.
استخدم صياغة قانونية دقيقة ومواد مرقمة. أعد النص فقط بدون أي تعليق.""",
        "fields": ["البضاعة/الخدمة الموردة", "الكمية الإجمالية", "سعر الوحدة", "إجمالي قيمة العقد", "مواعيد التسليم", "شروط الدفع"]
    },
    "عقد توظيف": {
        "system": """أنت محامٍ متخصص في قانون العمل وعقود التوظيف. اكتب عقد توظيف كامل ومفصل باللغة العربية الفصحى وفق قانون العمل في الدولة المحددة.
العقد يجب أن يتضمن: ديباجة، المسمى الوظيفي والمهام، الراتب والمزايا، ساعات العمل والإجازات، فترة التجربة، السرية وحظر المنافسة، إنهاء العقد ومكافأة نهاية الخدمة، الانضباط والجزاءات، فض النزاعات، أحكام ختامية، توقيعات.
استخدم صياغة قانونية دقيقة ومواد مرقمة. أعد النص فقط بدون أي تعليق.""",
        "fields": ["المسمى الوظيفي", "الراتب الأساسي الشهري", "مدة العقد (محدد/غير محدد)", "ساعات العمل الأسبوعية", "فترة التجربة", "بدلات ومزايا إضافية"]
    },
    "عقد مرابحة": {
        "system": """أنت محامٍ ومستشار شرعي متخصص في عقود المرابحة الإسلامية. اكتب عقد مرابحة كامل ومفصل باللغة العربية الفصحى وفق أحكام الشريعة الإسلامية والقانون المحدد.
العقد يجب أن يتضمن: ديباجة، وصف السلعة، ثمن الشراء وهامش الربح، إجمالي ثمن المرابحة، جدول السداد، ضمانات الوفاء، أحكام التأخر في السداد (وفق الضوابط الشرعية)، الرقابة الشرعية، فض النزاعات، أحكام ختامية، توقيعات.
استخدم صياغة قانونية وشرعية دقيقة ومواد مرقمة. أعد النص فقط بدون أي تعليق.""",
        "fields": ["وصف السلعة", "ثمن الشراء الأصلي", "هامش الربح (%)", "إجمالي ثمن المرابحة", "عدد الأقساط", "قيمة القسط الشهري"]
    }
}

def build_prompt(contract_type, jurisdiction, party1, party2, fields_data, extra_details):
    """Build a detailed prompt for the AI."""
    today = datetime.date.today()
    hijri_note = ""  # Could add hijri conversion if needed
    
    fields_text = "\n".join([f"- {k}: {v}" for k, v in fields_data.items() if v])
    
    prompt = f"""اكتب {contract_type} كاملاً وفق القانون في {jurisdiction}.

التاريخ: {today.strftime('%Y/%m/%d')}
الطرف الأول: {party1}
الطرف الثاني: {party2}

البيانات التفصيلية:
{fields_text}

{f'تفاصيل إضافية: {extra_details}' if extra_details else ''}

المتطلبات:
- اكتب العقد كاملاً بجميع موادّه وبنوده
- استخدم الأرقام العربية (١، ٢، ٣...)
- ابدأ كل مادة بـ "المادة الأولى:" "المادة الثانية:" إلخ
- أضف خط التوقيع في النهاية لكلا الطرفين
- لا تضف أي تعليق أو ملاحظة خارج نص العقد"""
    
    return prompt


def create_word_document(contract_text, contract_type, party1, party2, jurisdiction):
    """Create a properly formatted Word document from the contract text."""
    doc = Document()
    
    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    # ── RTL document direction ──
    doc.core_properties.language = 'ar-JO'
    
    # Helper: set paragraph RTL
    def set_rtl(para):
        pPr = para._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.insert(0, bidi)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
    
    # ── Header: Logo area ──
    header_para = doc.add_paragraph()
    set_rtl(header_para)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run('ميزان | LEGAL AI')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x4A)
    run.font.bold = True
    
    # ── Divider ──
    div_para = doc.add_paragraph('─' * 60)
    div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_para.runs[0].font.size = Pt(8)
    div_para.runs[0].font.color.rgb = RGBColor(0x88, 0x99, 0xAA)
    
    # ── Title ──
    title_para = doc.add_paragraph()
    set_rtl(title_para)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(contract_type)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x4A)
    
    # ── Subtitle: parties + jurisdiction ──
    sub_para = doc.add_paragraph()
    set_rtl(sub_para)
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f'{party1}  ←→  {party2}  |  {jurisdiction}')
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x44, 0x66, 0x88)
    
    doc.add_paragraph()  # spacer
    
    # ── Parse and render contract text ──
    lines = contract_text.strip().split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Detect article headers: المادة الأولى / المادة (١) / المادة 1
        is_article = bool(re.match(r'^المادة\s+(الأولى|الثانية|الثالثة|الرابعة|الخامسة|السادسة|السابعة|الثامنة|التاسعة|العاشرة|الحادية|الثانية عشرة|\(?\d+\)?|[١٢٣٤٥٦٧٨٩٠]+)', line))
        is_heading = line.startswith('##') or (len(line) < 60 and line.endswith(':') and not line.startswith('-'))
        
        if is_article:
            p = doc.add_paragraph()
            set_rtl(p)
            # Extract article title and body
            colon_idx = line.find(':')
            if colon_idx > 0:
                title_part = line[:colon_idx + 1]
                body_part = line[colon_idx + 1:].strip()
                run_title = p.add_run(title_part + ' ')
                run_title.font.bold = True
                run_title.font.size = Pt(12)
                run_title.font.color.rgb = RGBColor(0x1A, 0x2A, 0x4A)
                if body_part:
                    run_body = p.add_run(body_part)
                    run_body.font.size = Pt(11)
            else:
                run = p.add_run(line)
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x4A)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif is_heading:
            p = doc.add_paragraph()
            set_rtl(p)
            run = p.add_run(line.lstrip('#').strip())
            run.font.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1A, 0x2A, 0x4A)
            p.paragraph_format.space_before = Pt(10)
        elif line.startswith('-') or line.startswith('•') or line.startswith('*'):
            p = doc.add_paragraph(style='List Bullet')
            set_rtl(p)
            run = p.add_run(line.lstrip('-•* ').strip())
            run.font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            set_rtl(p)
            run = p.add_run(line)
            run.font.size = Pt(11)
        
        p.paragraph_format.line_spacing = Pt(18)
    
    # ── Signature block ──
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    set_rtl(sig_para)
    sig_run = sig_para.add_run('─' * 60)
    sig_run.font.size = Pt(8)
    sig_run.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)
    
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Table Grid'
    
    # Row 0: headers
    sig_table.cell(0, 0).text = f'الطرف الأول\n{party1}'
    sig_table.cell(0, 1).text = f'الطرف الثاني\n{party2}'
    # Row 1: signature line
    sig_table.cell(1, 0).text = 'التوقيع: ___________________'
    sig_table.cell(1, 1).text = 'التوقيع: ___________________'
    # Row 2: date
    sig_table.cell(2, 0).text = 'التاريخ: ___________________'
    sig_table.cell(2, 1).text = 'التاريخ: ___________________'
    
    for row in sig_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                set_rtl(para)
                for run in para.runs:
                    run.font.size = Pt(11)
    
    # ── Footer ──
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    set_rtl(footer_para)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f'تم إنشاء هذا العقد بواسطة ميزان Legal AI  |  {datetime.date.today().strftime("%Y/%m/%d")}')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0xAA, 0xBB)
    
    # ── Save to buffer ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route('/api/generate-contract', methods=['POST'])
def generate_contract():
    try:
        data = request.get_json()
        contract_type = data.get('contract_type', 'عقد إيجار تجاري')
        jurisdiction = data.get('jurisdiction', 'الأردن')
        party1 = data.get('party1', 'الطرف الأول')
        party2 = data.get('party2', 'الطرف الثاني')
        fields_data = data.get('fields', {})
        extra_details = data.get('extra_details', '')
        
        config = CONTRACT_CONFIGS.get(contract_type, CONTRACT_CONFIGS['عقد إيجار تجاري'])
        
        # Build prompt
        prompt = build_prompt(contract_type, jurisdiction, party1, party2, fields_data, extra_details)
        
        # Call OpenAI
        response = client.chat.completions.create(
            model='gpt-4.1-mini',
            messages=[
                {'role': 'system', 'content': config['system']},
                {'role': 'user', 'content': prompt}
            ],
            max_tokens=4000,
            temperature=0.3
        )
        
        contract_text = response.choices[0].message.content.strip()
        
        # Generate Word document
        doc_buffer = create_word_document(contract_text, contract_type, party1, party2, jurisdiction)
        
        # Create safe filename
        safe_type = contract_type.replace(' ', '_')
        filename = f'{safe_type}_{party1[:10]}_{datetime.date.today().strftime("%Y%m%d")}.docx'
        
        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/contract-fields', methods=['GET'])
def get_contract_fields():
    """Return the dynamic fields for a given contract type."""
    contract_type = request.args.get('type', 'عقد إيجار تجاري')
    config = CONTRACT_CONFIGS.get(contract_type, CONTRACT_CONFIGS['عقد إيجار تجاري'])
    return jsonify({'fields': config['fields']})


@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'service': 'meezan-contract-server'})


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5051, debug=False)
